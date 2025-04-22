from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .base import DocumentGeneratorBase

class WordGenerator(DocumentGeneratorBase):
    def create(self):
        """Word belgesi oluşturur"""
        doc = Document()
        # ... Word spesifik metodlar ...
        self._add_qr_to_doc(doc)  # QR kod sayfası ekle
        return self._save_document(doc)

    def _add_footer(self, doc, text):
        """Word belgesi için footer ekler"""
        for section in doc.sections:
            footer = section.footer
            p = footer.paragraphs[0]
            p.text = text
            # ... footer stili ...

    def _add_qr_to_doc(self, doc):
        """Word belgesine QR kod sayfası ekler"""
        qr_data = self._add_qr_page()
        
        doc.add_page_break()
        heading = doc.add_heading(qr_data['title'], level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(qr_data['qr_image'], width=Inches(2))
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(qr_data['caption'])
        run.font.size = Pt(11)
        run.font.name = 'Arial'

    # ... Diğer Word spesifik metodlar ...
