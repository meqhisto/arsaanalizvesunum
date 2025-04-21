from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib.pagesizes import A4
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os
import json
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor, Inches, Pt, Cm
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from datetime import datetime

class DocumentGenerator:
    def __init__(self, arsa_data, analiz_ozeti, file_id, output_dir):
        self.arsa_data = arsa_data
        self.analiz_ozeti = analiz_ozeti
        self.file_id = file_id
        self.output_dir = output_dir
        self.sunum_klasoru = os.path.join(output_dir, file_id)
        
        if not os.path.exists(self.sunum_klasoru):
            os.makedirs(self.sunum_klasoru)

    def _format_currency(self, value):
        """Para birimini formatlayan yardımcı metod"""
        try:
            return f"{float(value):,.2f} TL"
        except (TypeError, ValueError):
            return "0,00 TL"

    def _format_area(self, value):
        """Alan birimini formatlayan yardımcı metod"""
        try:
            return f"{float(value):,.2f} m²"
        except (TypeError, ValueError):
            return "0,00 m²"

    def _get_arsa_bilgileri(self):
        """Arsa bilgilerini hazırlayan yardımcı metod"""
        return [
            ['İl/İlçe', f"{self.arsa_data.get('il')}/{self.arsa_data.get('ilce')}"],
            ['Mahalle', self.arsa_data.get('mahalle', '')],
            ['Ada/Parsel', f"{self.arsa_data.get('ada')}/{self.arsa_data.get('parsel')}"],
            ['Alan', self._format_area(self.arsa_data.get('metrekare'))],
            ['İmar Durumu', self.arsa_data.get('imar_durumu', '')],
            ['TAKS/KAKS', f"{self.arsa_data.get('taks')}/{self.arsa_data.get('kaks')}"],
            ['Toplam Fiyat', self._format_currency(self.arsa_data.get('fiyat'))],
            ['m² Fiyatı', self._format_currency(float(self.arsa_data.get('fiyat', 0))/float(self.arsa_data.get('metrekare', 1)))],
        ]

    def _get_altyapi_durumu(self):
        """Altyapı durumunu hazırlayan yardımcı metod"""
        altyapi_list = self.arsa_data.get('altyapi[]', [])
        return [
            ['Yol', '✓' if 'yol' in altyapi_list else '✗'],
            ['Elektrik', '✓' if 'elektrik' in altyapi_list else '✗'],
            ['Su', '✓' if 'su' in altyapi_list else '✗'],
            ['Doğalgaz', '✓' if 'dogalgaz' in altyapi_list else '✗'],
            ['Kanalizasyon', '✓' if 'kanalizasyon' in altyapi_list else '✗']
        ]

    def _get_swot_analizi(self):
        """SWOT analizini hazırlayan yardımcı metod"""
        swot_data = {
            'Güçlü Yönler': self.arsa_data.get('strengths', []),
            'Zayıf Yönler': self.arsa_data.get('weaknesses', []),
            'Fırsatlar': self.arsa_data.get('opportunities', []),
            'Tehditler': self.arsa_data.get('threats', [])
        }
        
        # String ise JSON'a çevir
        for key, value in swot_data.items():
            if isinstance(value, str):
                try:
                    swot_data[key] = json.loads(value)
                except json.JSONDecodeError:
                    swot_data[key] = [value] if value else []
                    
        return swot_data

    def _set_cell_background(self, cell, color):
        """Tablo hücresinin arka plan rengini ayarlar"""
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def create_word(self):
        """Word belgesi oluşturur"""
        doc = Document()
        
        # Sayfa yapılandırması
        sections = doc.sections
        for section in sections:
            section.page_width = Inches(11.69)  # A4 genişlik
            section.page_height = Inches(8.27)  # A4 yükseklik - yatay görünüm
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
        
        # Kapak Sayfası
        title = doc.add_heading('', 0)
        title_run = title.add_run('ARSA ANALİZ RAPORU')
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(40)
        title_run.font.color.rgb = RGBColor(26, 35, 126) # Koyu Mavi
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Alt başlık
        subtitle = doc.add_paragraph() # Heading yerine paragraph kullanalım
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(f"{self.arsa_data.get('il', '')}, {self.arsa_data.get('ilce', '')}")
        subtitle_run.font.name = 'Arial'
        subtitle_run.font.size = Pt(28)
        subtitle_run.font.color.rgb = RGBColor(63, 81, 181) # Orta Mavi
        subtitle.space_before = Pt(12)
        subtitle.space_after = Pt(30)
        
        # Tarih
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_paragraph.add_run(datetime.now().strftime('%d.%m.%Y'))
        date_run.font.name = 'Arial'
        date_run.font.size = Pt(14)
        date_run.font.color.rgb = RGBColor(96, 125, 139) # Gri Mavi
        
        doc.add_page_break()
        
        # İçindekiler Sayfası
        doc.add_heading('İçindekiler', level=1).runs[0].font.size = Pt(24)
        sections = [
            'Arsa Bilgileri',
            'Altyapı Durumu',
            'SWOT Analizi',
            'Analiz Özeti'
        ]
        
        for i, section in enumerate(sections, 1):
            p = doc.add_paragraph()
            p.add_run(f"{i}. {section}").font.size = Pt(14)
            
        doc.add_page_break()
        
        # Arsa Bilgileri Sayfası
        heading = doc.add_heading('Arsa Bilgileri', 1)
        heading.runs[0].font.size = Pt(24)
        heading.runs[0].font.color.rgb = RGBColor(26, 35, 126)
        
        # Info box - Lokasyon
        info_table = doc.add_table(rows=1, cols=1)
        info_table.style = 'Light Grid'
        cell = info_table.rows[0].cells[0]
        self._set_cell_background(cell, "F5F5F5")
        
        p = cell.paragraphs[0]
        p.add_run('📍 Lokasyon\n').bold = True
        p.add_run(f"\n{self.arsa_data.get('il')}/{self.arsa_data.get('ilce')}\n")
        p.add_run(f"Mahalle: {self.arsa_data.get('mahalle')}\n")
        p.add_run(f"Ada/Parsel: {self.arsa_data.get('ada')}/{self.arsa_data.get('parsel')}")
        p.alignment= WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Ana bilgi tablosu - 3 sütunlu modern tasarım
        data = [
            ['Alan', 'İmar Durumu', 'Fiyat'],
            [
                self._format_area(self.arsa_data.get('metrekare')),
                self.arsa_data.get('imar_durumu', ''),
                self._format_currency(self.arsa_data.get('fiyat'))
            ],
            ['TAKS', 'KAKS', 'm² Fiyatı'],
            [
                str(self.arsa_data.get('taks')),
                str(self.arsa_data.get('kaks')),
                self._format_currency(float(self.arsa_data.get('fiyat', 0))/float(self.arsa_data.get('metrekare', 1)))
            ]
        ]
        
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for i, row in enumerate(data):
            for j, val in enumerate(row):
                cell = table.cell(i, j)
                cell.text = val
                if i % 2 == 0:  # Başlık satırları
                    self._set_cell_background(cell, "1A237E")
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.bold = True
                else:  # Veri satırları
                    self._set_cell_background(cell, "E8EAF6")
        
    
        
        # Altyapı Durumu - Modern gösterim
        heading = doc.add_heading('Altyapı Durumu', 1)
        heading.runs[0].font.size = Pt(24)
        heading.runs[0].font.color.rgb = RGBColor(26, 35, 126)
        
        altyapi_data = self._get_altyapi_durumu()
        table = doc.add_table(rows=1, cols=len(altyapi_data))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # İkon ve durum gösterimi
        for i, (altyapi, durum) in enumerate(altyapi_data):
            cell = table.cell(0, i)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # İkon ve başlık
            run = paragraph.add_run(f'{altyapi}\n\n')
            run.font.bold = True
            
            # Durum ikonu
            status_run = paragraph.add_run('✓' if durum == '✓' else '✗')
            status_run.font.size = Pt(24)
            status_run.font.color.rgb = RGBColor(76, 175, 80) if durum == '✓' else RGBColor(244, 67, 54)
            
            # Arka plan rengi
            self._set_cell_background(cell, "E8EAF6")
            
       

        # SWOT Analizi - Modern 2x2 grid
        swot_heading = doc.add_heading('SWOT Analizi', 1)
        swot_heading.runs[0].font.color.rgb = RGBColor(26, 35, 126)
        
        # SWOT verilerini 2x2 tablo olarak göster
        swot_table = doc.add_table(rows=2, cols=2)
        swot_table.style = 'Table Grid'
        swot_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        swot_data = self._get_swot_analizi()
        cell_data = [
            ('Güçlü Yönler', swot_data['Güçlü Yönler'], "4CAF50"),  # Yeşil
            ('Zayıf Yönler', swot_data['Zayıf Yönler'], "F44336"),  # Kırmızı
            ('Fırsatlar', swot_data['Fırsatlar'], "2196F3"),        # Mavi
            ('Tehditler', swot_data['Tehditler'], "FF9800")         # Turuncu
        ]
        
        for i, (title, items, color) in enumerate(cell_data):
            cell = swot_table.cell(i//2, i%2)
            self._set_cell_background(cell, color)
            
            # Başlık
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
            # Maddeler
            for item in items:
                p = cell.add_paragraph()
                p.style = 'List Bullet'
                run = p.add_run(item)
                run.font.color.rgb = RGBColor(255, 255, 255)
        
        doc.add_page_break()
        
        # Analiz Özeti - Modern tasarım
        ozet_heading = doc.add_heading('Analiz Özeti', 1)
        ozet_heading.runs[0].font.color.rgb = RGBColor(26, 35, 126)
        ozet_heading.runs[0].font.size = Pt(24)
        
        # Temel Özet
        temel_ozet = doc.add_paragraph()
        temel_ozet.add_run('Temel Değerlendirme\n').bold = True
        temel_ozet.add_run(self.analiz_ozeti.get('temel_ozet', ''))
        
        doc.add_paragraph()  # Boşluk
        
        # Yatırım Özeti
        yatirim_ozet = doc.add_paragraph()
        yatirim_ozet.add_run('Yatırım Değerlendirmesi\n').bold = True
        yatirim_ozet.add_run(self.analiz_ozeti.get('yatirim_ozet', ''))
        
        doc.add_paragraph()  # Boşluk
        
        # Tavsiyeler
        tavsiyeler = doc.add_paragraph()
        tavsiyeler.add_run('Öneriler ve Tavsiyeler\n').bold = True
        tavsiyeler.add_run(self.analiz_ozeti.get('tavsiyeler', ''))
        
        filename = os.path.join(self.sunum_klasoru, f'analiz_{self.file_id}.docx')
        doc.save(filename)
        return filename

    def _register_fonts(self):
        """Fontları kaydet"""
        WINDOWS_FONTS = os.path.join(os.environ['SYSTEMROOT'], 'Fonts')
        fonts = {
            'arial': 'arial.ttf',
            'arial-bold': 'arialbd.ttf',
            'tahoma': 'tahoma.ttf',
            'tahoma-bold': 'tahomabd.ttf'
        }
        
        registered_fonts = {}
        for font_name, font_file in fonts.items():
            try:
                font_path = os.path.join(WINDOWS_FONTS, font_file)
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
                    registered_fonts[font_name] = True
                    continue
            except:
                pass
            registered_fonts[font_name] = False
        
        # En az bir font bulunduysa onu kullan, bulunamadıysa Helvetica
        if registered_fonts.get('arial'):
            return 'arial', 'arial-bold'
        elif registered_fonts.get('tahoma'):
            return 'tahoma', 'tahoma-bold'
        else:
            return 'Helvetica', 'Helvetica-Bold'

    def create_pdf(self):
        """PDF belgesi oluşturur"""
        # Font ayarları
        base_font, base_font_bold = self._register_fonts()
        
        filename = os.path.join(self.sunum_klasoru, f'analiz_{self.file_id}.pdf')
        doc = SimpleDocTemplate(
            filename,
            pagesize=A4,
            leftMargin=1.5*cm,
            rightMargin=1.5*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        
        # Stil tanımlamaları güncellendi
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=base_font_bold,
            fontSize=24,
            alignment=TA_CENTER,
            spaceAfter=30
        ))
        
        styles.add(ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontName=base_font_bold,
            fontSize=16,
            spaceBefore=20,
            spaceAfter=15,
            encoding='utf-8'  # UTF-8 encoding eklendi
        ))
        
        styles.add(ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName=base_font,
            fontSize=11,
            spaceAfter=8,
            encoding='utf-8',  # UTF-8 encoding eklendi
            leading=14
        ))
        
        # Tablo stili
        table_style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a237e')),  # Koyu mavi
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), base_font_bold),
            ('FONTNAME', (0, 1), (-1, -1), base_font),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('FONTSIZE', (0, 1), (-1, -1), 11),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#e8eaf6')),  # Açık mavi
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#3949ab')),  # Orta mavi
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ])

        elements = []
        
        # Başlık
        elements.append(Paragraph('Arsa Analiz Raporu', styles['CustomTitle']))
        elements.append(Paragraph(f"{self.arsa_data.get('il')}, {self.arsa_data.get('ilce')}", styles['CustomHeading']))
        elements.append(Spacer(1, 20))
        
        # Arsa Bilgileri Tablosu
        elements.append(Paragraph('Arsa Bilgileri', styles['CustomHeading']))
        data = [['Parsel Bilgileri', 'Açıklama']] + self._get_arsa_bilgileri()
        table = Table(data, colWidths=[doc.width*0.3, doc.width*0.7])
        table.setStyle(table_style)
        elements.append(table)
        elements.append(Spacer(1, 20))
        
        # Altyapı Durumu
        elements.append(Paragraph('Altyapı Durumu', styles['CustomHeading']))
        altyapi = self._get_altyapi_durumu()
        # Modern kutu dolgu: ✓ ise kutu yeşil, ✗ ise kutu kırmızı
        altyapi_table_data = []
        for ozellik, durum in altyapi:
            color = colors.HexColor('#4CAF50') if durum == '✓' else colors.HexColor('#F44336')
            # Kutuya sadece ✓ veya ✗ ortalanmış şekilde
            altyapi_table_data.append([
            Paragraph(ozellik, styles['CustomNormal']),
            Table(
                [[Paragraph(durum, styles['CustomNormal'])]],
                colWidths=[doc.width*0.1],
                rowHeights=[18],
                style=TableStyle([
                ('BACKGROUND', (0, 0), (0, 0), color),
                ('ALIGN', (0, 0), (0, 0), 'CENTER'),
                ('VALIGN', (0, 0), (0, 0), 'MIDDLE'),
                ('TEXTCOLOR', (0, 0), (0, 0), colors.whitesmoke),
                ('FONTSIZE', (0, 0), (0, 0), 14),
                ('BOX', (0, 0), (0, 0), 1, color),
                ])
            )
            ])
        data = [['Özellik', 'Durum']] + altyapi_table_data
        table = Table(data, colWidths=[doc.width*0.3, doc.width*0.7])
        table.setStyle(table_style)
        elements.append(table)
        elements.append(Spacer(1, 20))
        
        # SWOT Analizi
        elements.append(PageBreak())
        elements.append(Paragraph('SWOT Analizi', styles['CustomHeading']))

        # Prepare SWOT data for 2x2 table
        swot = self._get_swot_analizi()
        swot_titles = ['Güçlü Yönler', 'Zayıf Yönler', 'Fırsatlar', 'Tehditler']
        swot_colors = ['#4CAF50', '#F44336', '#2196F3', '#FF9800']  # Green, Red, Blue, Orange

        # Prepare cell content as Paragraphs with bullet points
        swot_cells = []
        for idx, title in enumerate(swot_titles):
            items = swot.get(title, [])
            cell_content = [Paragraph(f'<b>{title}</b>', styles['CustomNormal'])]
            for item in items:
                cell_content.append(Paragraph(f'• {item}', styles['CustomNormal']))
            swot_cells.append(cell_content)

        # 2x2 grid
        data = [
            [swot_cells[0], swot_cells[1]],
            [swot_cells[2], swot_cells[3]],
        ]

        swot_table = Table(data, colWidths=[doc.width/2, doc.width/2])
        swot_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, 0), colors.HexColor(swot_colors[0])),
            ('BACKGROUND', (1, 0), (1, 0), colors.HexColor(swot_colors[1])),
            ('BACKGROUND', (0, 1), (0, 1), colors.HexColor(swot_colors[2])),
            ('BACKGROUND', (1, 1), (1, 1), colors.HexColor(swot_colors[3])),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.whitesmoke),
            ('INNERGRID', (0, 0), (-1, -1), 1, colors.white),
            ('BOX', (0, 0), (-1, -1), 2, colors.HexColor('#3949ab')),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ]))
        elements.append(swot_table)
        elements.append(Spacer(1, 20))
            
        # Analiz Özeti
        elements.append(PageBreak())
        elements.append(Paragraph('Analiz Özeti', styles['CustomHeading']))
        elements.append(Paragraph(self.analiz_ozeti.get('temel_ozet', ''), styles['CustomNormal']))
        elements.append(Spacer(1, 10))
        elements.append(Paragraph(self.analiz_ozeti.get('yatirim_ozet', ''), styles['CustomNormal']))
        elements.append(Spacer(1, 10))
        elements.append(Paragraph(self.analiz_ozeti.get('tavsiyeler', ''), styles['CustomNormal']))
        
        doc.build(elements)
        return filename
